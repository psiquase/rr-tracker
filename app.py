"""
╔══════════════════════════════════════════════════════════╗
║   PRODUÇÃO PHOENIX  ◆  WEB EDITION                   ║
║   Flask + SQLite  |  Free hosting via Render + GitHub     ║
╚══════════════════════════════════════════════════════════╝
"""

import os
import io
import json
import sqlite3
from functools import wraps
from datetime import datetime, timedelta, date
from flask import (Flask, render_template, request, jsonify,
                   redirect, url_for, session, send_file)
app = Flask(__name__)

# ── Secret key for sessions (change this in production!) ─────
app.secret_key = os.environ.get("SECRET_KEY", "phoenix-pixel-2026")

# ── Session config ────────────────────────────────────────────
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(hours=12)

# ── Admin password (set via env var ADMIN_PASSWORD on Render) ─
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "prota2026")

# ── Database path ─────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH  = os.path.join(BASE_DIR, "data", "rerecording.db")


# ── Inject is_admin into every template ──────────────────────
@app.context_processor
def inject_auth():
    return dict(is_admin=is_admin())


# ═══════════════════════════════════════════
#  AUTH HELPERS
# ═══════════════════════════════════════════

def is_admin():
    return session.get("admin") is True

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not is_admin():
            return redirect(url_for("login", next=request.path))
        return f(*args, **kwargs)
    return decorated


# ═══════════════════════════════════════════
#  DATABASE HELPERS
# ═══════════════════════════════════════════

def get_db():
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    with get_db() as conn:
        conn.executescript("""
            CREATE TABLE IF NOT EXISTS records (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                producer    TEXT    NOT NULL,
                project     TEXT    NOT NULL,
                count       INTEGER NOT NULL DEFAULT 1,
                reason      TEXT    NOT NULL,
                description TEXT    DEFAULT '',
                observation TEXT    DEFAULT '',
                created_at  TEXT    NOT NULL
            );

            CREATE TABLE IF NOT EXISTS productions (
                id           INTEGER PRIMARY KEY AUTOINCREMENT,
                title        TEXT    NOT NULL,
                producer     TEXT    NOT NULL,
                prod_type    TEXT    DEFAULT 'producao',
                total_arcs   INTEGER NOT NULL DEFAULT 11,
                arcs_done    TEXT    DEFAULT '[]',
                status       TEXT    DEFAULT 'iniciado',
                started_at   TEXT    NOT NULL,
                updated_at   TEXT,
                paused_at    TEXT,
                completed_at TEXT,
                notes        TEXT    DEFAULT '',
                script_chars INTEGER DEFAULT 0,
                arc_chars    TEXT    DEFAULT ''
            );

            CREATE TABLE IF NOT EXISTS production_daily (
                id            INTEGER PRIMARY KEY AUTOINCREMENT,
                production_id INTEGER NOT NULL,
                log_date      TEXT    NOT NULL,
                chars_written INTEGER DEFAULT 0,
                notes         TEXT    DEFAULT '',
                FOREIGN KEY (production_id) REFERENCES productions(id)
            );
        """)

        # ── Shorts table (new clean schema) ──────────────────
        # Migrate if old schema exists (had 'chars' column instead of started_at)
        cols = [r[1] for r in conn.execute("PRAGMA table_info(shorts)").fetchall()]
        if cols and 'started_at' not in cols:
            conn.execute("DROP TABLE shorts")
            conn.commit()

        conn.executescript("""
            CREATE TABLE IF NOT EXISTS shorts (
                id           INTEGER PRIMARY KEY AUTOINCREMENT,
                title        TEXT    NOT NULL,
                producer     TEXT    NOT NULL,
                status       TEXT    DEFAULT 'iniciado',
                started_at   TEXT    NOT NULL,
                completed_at TEXT,
                paused_at    TEXT,
                notes        TEXT    DEFAULT ''
            );
        """)

        # ── Productions: migrate new columns ─────────────────
        prod_cols = [r[1] for r in conn.execute("PRAGMA table_info(productions)").fetchall()]
        if 'script_chars' not in prod_cols:
            conn.execute("ALTER TABLE productions ADD COLUMN script_chars INTEGER DEFAULT 0")
        if 'arc_chars' not in prod_cols:
            conn.execute("ALTER TABLE productions ADD COLUMN arc_chars TEXT DEFAULT ''")
        if 'prod_type' not in prod_cols:
            conn.execute("ALTER TABLE productions ADD COLUMN prod_type TEXT DEFAULT 'producao'")
        conn.commit()


init_db()


def db_stats():
    with get_db() as c:
        total   = c.execute("SELECT COALESCE(SUM(count),0) FROM records").fetchone()[0]
        top_p   = c.execute(
            "SELECT producer, SUM(count) t FROM records GROUP BY producer ORDER BY t DESC LIMIT 1"
        ).fetchone()
        top_r   = c.execute(
            "SELECT reason, SUM(count) t FROM records GROUP BY reason ORDER BY t DESC LIMIT 1"
        ).fetchone()
        week_t  = c.execute(
            "SELECT COALESCE(SUM(count),0) FROM records WHERE date(created_at)>=date('now','weekday 0','-7 days')"
        ).fetchone()[0]
    return {
        "total":        total,
        "week_total":   week_t,
        "top_producer": dict(top_p) if top_p else None,
        "top_reason":   dict(top_r) if top_r else None,
    }


def db_reason_dist():
    with get_db() as c:
        return [dict(r) for r in c.execute(
            "SELECT reason, SUM(count) total FROM records GROUP BY reason ORDER BY total DESC"
        ).fetchall()]


def db_by_producer():
    with get_db() as c:
        return [dict(r) for r in c.execute(
            """SELECT producer, SUM(count) total,
               (SELECT reason FROM records r2
                WHERE r2.producer=r.producer
                GROUP BY reason ORDER BY SUM(count) DESC LIMIT 1) top_reason
               FROM records r GROUP BY producer ORDER BY total DESC"""
        ).fetchall()]


def db_weekly(weeks_back=0):
    today = date.today()
    mon   = today - timedelta(days=today.weekday()) - timedelta(weeks=weeks_back)
    sun   = mon + timedelta(days=6)
    with get_db() as c:
        rows = c.execute(
            """SELECT * FROM records
               WHERE date(created_at) BETWEEN ? AND ?
               ORDER BY created_at DESC""",
            (mon.isoformat(), sun.isoformat())
        ).fetchall()
    return [dict(r) for r in rows], mon, sun


def db_trend():
    with get_db() as c:
        rows = c.execute(
            """SELECT strftime('%d/%m', created_at) day, SUM(count) total
               FROM records
               WHERE date(created_at) >= date('now', '-14 days')
               GROUP BY day ORDER BY day"""
        ).fetchall()
    return [dict(r) for r in rows]


def calc_bdays(start_str, end_str=None):
    # Business days between two date strings.
    try:
        start = date.fromisoformat(str(start_str)[:10])
        end   = date.fromisoformat(str(end_str)[:10]) if end_str else date.today()
    except Exception:
        return 0
    if start > end:
        return 0
    count = 0
    cur = start
    while cur < end:
        if cur.weekday() < 5:
            count += 1
        cur += timedelta(days=1)
    return count


def db_production_report():
    # Returns per-producer summary + all productions grouped by type
    with get_db() as c:
        all_prods = [dict(r) for r in c.execute(
            "SELECT * FROM productions ORDER BY producer, started_at DESC"
        ).fetchall()]

    # Enrich each production
    for p in all_prods:
        p['prod_type']       = p.get('prod_type') or 'producao'
        p['arcs_done_list']  = json.loads(p.get('arcs_done') or '[]')
        p['arcs_done_count'] = len(p['arcs_done_list'])
        if p['prod_type'] == 'producao':
            if p['status'] == 'concluido' and p.get('completed_at'):
                p['duration_bdays'] = calc_bdays(p['started_at'], p['completed_at'])
            else:
                p['duration_bdays'] = calc_bdays(p['started_at'])
            p['duration_label'] = f"{p['duration_bdays']} d.u."
            p['duration_minutes'] = 0
        else:
            if p['status'] == 'concluido' and p.get('completed_at'):
                p['duration_minutes'] = calc_minutes(p['started_at'], p['completed_at'])
            else:
                p['duration_minutes'] = calc_minutes(p['started_at'], datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            p['duration_label']   = format_duration(p['duration_minutes'])
            p['duration_bdays']   = 0

    # Per-type summaries
    type_labels = {'producao':'PRODUÇÃO','roteiro':'ROTEIRO','decupagem':'DECUPAGEM','edicao':'EDIÇÃO'}
    types_summ = {}
    for type_id, label in type_labels.items():
        items = [p for p in all_prods if p['prod_type']==type_id]
        concluidos = [p for p in items if p['status']=='concluido']
        if type_id == 'producao':
            avg = round(sum(p['duration_bdays'] for p in concluidos)/len(concluidos),1) if concluidos else None
            avg_label = f"{avg} d.u." if avg is not None else None
        else:
            avg = round(sum(p['duration_minutes'] for p in concluidos)/len(concluidos)) if concluidos else None
            avg_label = format_duration(avg) if avg is not None else None
        types_summ[type_id] = {
            'label':      label,
            'total':      len(items),
            'andamento':  sum(1 for p in items if p['status'] in ('iniciado','em_andamento')),
            'pausado':    sum(1 for p in items if p['status']=='pausado'),
            'concluido':  len(concluidos),
            'avg_label':  avg_label,
            'items':      items,
        }

    # Per-producer summary (producao only — days)
    producers = {}
    for p in all_prods:
        name = p['producer']
        if name not in producers:
            producers[name] = {
                'producer':    name, 'total': 0,
                'iniciado':    0, 'em_andamento': 0,
                'pausado':     0, 'concluido': 0,
                'avg_bdays':   None, '_durations': [],
            }
        g = producers[name]
        g['total'] += 1
        if p['status'] in g: g[p['status']] += 1
        if p['prod_type']=='producao' and p['status']=='concluido':
            g['_durations'].append(p['duration_bdays'])
    for g in producers.values():
        d = g['_durations']
        g['avg_bdays'] = round(sum(d)/len(d),1) if d else None
        del g['_durations']

    return list(producers.values()), all_prods, types_summ


# ═══════════════════════════════════════════
#  ROUTES
# ═══════════════════════════════════════════

@app.route("/")
def index():
    if is_admin():
        return redirect(url_for("dashboard"))
    return redirect(url_for("register"))


# ── AUTH ──────────────────────────────────

@app.route("/login", methods=["GET", "POST"])
def login():
    error    = None
    next_url = request.args.get("next", "/dashboard")
    if request.method == "POST":
        pwd      = request.form.get("password", "")
        next_url = request.form.get("next", "/dashboard")
        if pwd == ADMIN_PASSWORD:
            session.permanent = True
            session["admin"]  = True
            return redirect(next_url)
        error = "Senha incorreta. Tente novamente."
    return render_template("login.html", error=error,
                           next=next_url, active=None)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("register"))


# ── PROTECTED VIEWS ───────────────────────

@app.route("/dashboard")
@login_required
def dashboard():
    stats      = db_stats()
    reasons    = db_reason_dist()
    producers  = db_by_producer()
    trend      = db_trend()
    return render_template("dashboard.html",
                           stats=stats, reasons=reasons,
                           producers=producers, trend=trend,
                           active="dashboard")


@app.route("/register", methods=["GET", "POST"])
def register():
    msg = None
    msg_type = None
    if request.method == "POST":
        producer = request.form.get("producer", "").strip()
        project  = request.form.get("project",  "").strip()
        count_s  = request.form.get("count",    "1").strip()
        reason   = request.form.get("reason",   "")
        desc     = request.form.get("description", "").strip()
        obs      = request.form.get("observation",  "").strip()

        errors = []
        if not producer:        errors.append("Produtor obrigatório")
        if not project:         errors.append("Projeto obrigatório")
        if not count_s.isdigit() or int(count_s) < 1:
            errors.append("Número de regravaçôes inválido")
        if reason == "Outro" and not desc:
            errors.append("Descrição obrigatória quando motivo = 'Outro'")

        if errors:
            msg      = " | ".join(errors)
            msg_type = "error"
        else:
            now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            with get_db() as c:
                c.execute(
                    "INSERT INTO records (producer,project,count,reason,description,observation,created_at) VALUES (?,?,?,?,?,?,?)",
                    (producer, project, int(count_s), reason, desc, obs, now)
                )
            msg      = f"✔  Registro salvo: {producer} — {project} ({count_s}x)"
            msg_type = "success"

    reasons_list = [
        "Erro de roteiro", "Falta de atenção", "Erro técnico",
        "Não seguiu padrão", "Mudança de ideia (cliente)", "Outro",
    ]
    return render_template("register.html",
                           reasons=reasons_list,
                           msg=msg, msg_type=msg_type,
                           active="register")


@app.route("/reports")
@login_required
def reports():
    weeks_back   = int(request.args.get("w", 0))
    tab          = request.args.get("tab", "regravaçoes")
    rows, mon, sun = db_weekly(weeks_back)

    for r in rows:
        r["tag"] = ("high"   if r["count"] >= 5 else
                    "medium" if r["count"] >= 3 else "low")

    total = sum(r["count"] for r in rows)

    # Production report data
    prod_by_producer, all_prods, types_summ = db_production_report()
    shorts_summ, all_shorts     = db_shorts_report()

    return render_template("reports.html",
                           rows=rows, total=total,
                           mon=mon.strftime("%d/%m"),
                           sun=sun.strftime("%d/%m/%Y"),
                           weeks_back=weeks_back,
                           tab=tab,
                           prod_by_producer=prod_by_producer,
                           all_prods=all_prods,
                           types_summ=types_summ,
                           shorts_summ=shorts_summ,
                           all_shorts=[enrich_short(s) for s in all_shorts],
                           active="reports")


# ── JSON API (used by Chart.js calls) ────────────────────────
@app.route("/api/chart-data")
@login_required
def chart_data():
    return jsonify({
        "reasons":   db_reason_dist(),
        "producers": db_by_producer(),
        "trend":     db_trend(),
    })


# ── Health check ──────────────────────────────────────────────
@app.route("/health")
def health():
    return "OK", 200


# ═══════════════════════════════════════════
#  SHORTS HELPERS
# ═══════════════════════════════════════════

def db_shorts_all():
    with get_db() as c:
        return [dict(r) for r in c.execute(
            "SELECT * FROM shorts ORDER BY started_at DESC"
        ).fetchall()]


def db_shorts_report():
    # Average completion time per producer for shorts.
    rows = db_shorts_all()
    producers = {}
    for s in rows:
        name = s['producer']
        if name not in producers:
            producers[name] = {
                'producer':   name,
                'total':      0,
                'concluido':  0,
                'andamento':  0,
                'pausado':    0,
                'avg_bdays':  None,
                '_durs':      [],
            }
        g = producers[name]
        g['total'] += 1
        st = s['status']
        if st == 'concluido':
            g['concluido'] += 1
            if s.get('completed_at'):
                g['_durs'].append(calc_minutes(s['started_at'], s['completed_at']))
        elif st == 'pausado':
            g['pausado'] += 1
        else:
            g['andamento'] += 1
    for g in producers.values():
        if g['_durs']:
            avg_mins = round(sum(g['_durs'])/len(g['_durs']))
            g['avg_minutes'] = avg_mins
            g['avg_label']   = format_duration(avg_mins)
        else:
            g['avg_minutes'] = None
            g['avg_label']   = None
        del g['_durs']
    return list(producers.values()), rows


def enrich_short(s):
    d = dict(s)
    if d['status'] == 'concluido' and d.get('completed_at'):
        mins = calc_minutes(d['started_at'], d['completed_at'])
        d['duration_minutes'] = mins
        d['duration_label']   = format_duration(mins)
        d['duration_bdays']   = calc_bdays(d['started_at'], d['completed_at'])
    else:
        # In progress — show elapsed minutes
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        mins = calc_minutes(d['started_at'], now)
        d['duration_minutes'] = mins
        d['duration_label']   = format_duration(mins) + " (em curso)"
        d['duration_bdays']   = 0
    d['started_display']   = d['started_at'][:16].replace('T',' ')
    d['completed_display'] = d['completed_at'][:16].replace('T',' ') if d.get('completed_at') else '—'
    return d


# ═══════════════════════════════════════════
#  SHORTS ROUTES


@app.route("/shorts")
def shorts():
    all_s = [enrich_short(s) for s in db_shorts_all()]
    summ  = {
        'total':     len(all_s),
        'andamento': sum(1 for s in all_s if s['status'] in ('iniciado','em_andamento')),
        'concluido': sum(1 for s in all_s if s['status'] == 'concluido'),
        'pausado':   sum(1 for s in all_s if s['status'] == 'pausado'),
    }
    return redirect(url_for("productions"))


@app.route("/shorts/new", methods=["GET", "POST"])
def shorts_new():
    msg = msg_type = None
    from_prod = request.args.get("from") == "production_new"
    if request.method == "POST":
        title    = request.form.get("title",    "").strip()
        producer = request.form.get("producer", "").strip()
        notes    = request.form.get("notes",    "").strip()
        errors   = []
        if not title:    errors.append("Título obrigatório")
        if not producer: errors.append("Produtor obrigatório")
        if errors:
            msg = " | ".join(errors); msg_type = "error"
        else:
            now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            with get_db() as c:
                c.execute(
                    "INSERT INTO shorts (title,producer,status,started_at,notes) VALUES (?,?,?,?,?)",
                    (title, producer, "em_andamento", now, notes)
                )
            # If came from production_new page, show success there
            if request.form.get("from_prod") == "1" or from_prod:
                msg = f"✔  Short iniciado: {title}"; msg_type = "success"
            else:
                return redirect(url_for("productions"))
    return render_template("shorts_new.html", msg=msg, msg_type=msg_type,
                           active="production_new" if from_prod else "shorts")


@app.route("/shorts/<int:sid>/status", methods=["POST"])
def shorts_status(sid):
    new_status = request.form.get("status", "")
    if new_status not in ("em_andamento", "pausado", "concluido", "iniciado"):
        return redirect(url_for("productions"))
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with get_db() as c:
        if new_status == "concluido":
            c.execute("UPDATE shorts SET status=?, completed_at=? WHERE id=?",
                      (new_status, now, sid))
        elif new_status == "pausado":
            c.execute("UPDATE shorts SET status=?, paused_at=? WHERE id=?",
                      (new_status, now, sid))
        else:
            c.execute("UPDATE shorts SET status=? WHERE id=?", (new_status, sid))
    return redirect(url_for("productions"))


@app.route("/shorts/<int:sid>/delete", methods=["POST"])
@login_required
def shorts_delete(sid):
    with get_db() as c:
        c.execute("DELETE FROM shorts WHERE id=?", (sid,))
    return redirect(url_for("productions"))


@app.route("/records/<int:rid>/delete", methods=["POST"])
@login_required
def record_delete(rid):
    # Delete a re-recording record (admin only).
    with get_db() as c:
        c.execute("DELETE FROM records WHERE id=?", (rid,))
    # Return to reports page preserving tab and week
    ref = request.referrer or "/reports?tab=regravaçoes"
    return redirect(ref)


# ═══════════════════════════════════════════
#  PRODUCTION HELPERS
# ═══════════════════════════════════════════

def business_days_since(start_str):
    # Count business days from start to today, inclusive (day 1 = started today)
    try:
        start = date.fromisoformat(str(start_str)[:10])
    except Exception:
        return 1
    today = date.today()
    if start > today:
        return 1
    count = 0
    cur = start
    while cur <= today:
        if cur.weekday() < 5:  # Mon-Fri
            count += 1
        cur += timedelta(days=1)
    return max(1, count)  # minimum 1 (started today)


def calc_minutes(start_str, end_str):
    """Calculate minutes between two datetime strings."""
    try:
        from datetime import datetime as dt
        fmt = "%Y-%m-%d %H:%M:%S"
        start = dt.strptime(start_str[:19], fmt)
        end   = dt.strptime(end_str[:19], fmt)
        delta = end - start
        return max(0, int(delta.total_seconds() / 60))
    except Exception:
        return 0


def format_duration(minutes):
    """Format minutes into human-readable string."""
    if minutes < 60:
        return f"{minutes} min"
    h = minutes // 60
    m = minutes % 60
    if m == 0:
        return f"{h}h"
    return f"{h}h {m}min"


def deadline_color(bdays, status):
    if status == 'concluido':
        return 'green'
    if bdays <= 2:
        return 'green'
    if bdays == 3:
        return 'yellow'
    return 'red'






def prod_total_chars(prod_id):
    with get_db() as c:
        r = c.execute(
            "SELECT COALESCE(SUM(chars_written),0) FROM production_daily WHERE production_id=?",
            (prod_id,)
        ).fetchone()
    return r[0] if r else 0


def prod_today_chars(prod_id):
    today = date.today().isoformat()
    with get_db() as c:
        r = c.execute(
            "SELECT COALESCE(SUM(chars_written),0) FROM production_daily WHERE production_id=? AND log_date=?",
            (prod_id, today)
        ).fetchone()
    return r[0] if r else 0


def enrich_production(p):
    # Add computed fields to a production dict.
    d = dict(p)
    d['arcs_done_list']  = json.loads(d.get('arcs_done') or '[]')
    d['arcs_done_count'] = len(d['arcs_done_list'])
    d['bdays']           = business_days_since(d['started_at'])
    d['dl_color']        = deadline_color(d['bdays'], d['status'])
    d['prod_type']        = d.get('prod_type') or 'producao'
    d['script_chars']    = d.get('script_chars') or 0

    # Duration in minutes (for roteiro/decupagem/edicao — like Shorts)
    if d['prod_type'] != 'producao':
        if d['status'] == 'concluido' and d.get('completed_at'):
            d['duration_minutes'] = calc_minutes(d['started_at'], d['completed_at'])
        else:
            d['duration_minutes'] = calc_minutes(d['started_at'], datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        d['duration_label'] = format_duration(d['duration_minutes']) + ('' if d['status']=='concluido' else ' ⏳')
    else:
        d['duration_minutes'] = 0
        d['duration_label']   = f"{d.get('duration_bdays',0)} d.u."
    raw_arc              = d.get('arc_chars') or '{}'
    try:
        d['arc_chars_map'] = {int(k): v for k,v in json.loads(raw_arc).items()}
    except Exception:
        d['arc_chars_map'] = {}
    # Use sum of arc_chars_map as the authoritative total (FALAS [US] only)
    # Fall back to script_chars if no per-arc data exists
    arc_map_total = sum(d['arc_chars_map'].values()) if d['arc_chars_map'] else 0
    if arc_map_total > 0:
        d['script_chars'] = arc_map_total
    d['arc_chars_avg']   = round(d['script_chars'] / d['total_arcs']) if d['script_chars'] and d['total_arcs'] else 0
    d['arc_pct']         = int(d['arcs_done_count'] / d['total_arcs'] * 100) if d['total_arcs'] else 0
    d['total_chars']     = prod_total_chars(d['id'])
    d['today_chars']     = prod_today_chars(d['id'])
    goal = 5000 * 4  # 4 dias úteis × 5000 chars
    d['chars_pct']       = min(100, int(d['total_chars'] / goal * 100)) if goal else 0
    d['today_pct']       = min(100, int(d['today_chars'] / 5000 * 100))
    if d['status'] == 'concluido' and d.get('completed_at'):
        d['duration_bdays'] = calc_bdays(d['started_at'], d['completed_at'])
    else:
        d['duration_bdays'] = calc_bdays(d['started_at'])
    return d


# ═══════════════════════════════════════════
#  PRODUCTION ROUTES
# ═══════════════════════════════════════════

@app.route("/productions")
def productions():
    with get_db() as c:
        rows = c.execute("SELECT * FROM productions ORDER BY started_at DESC").fetchall()
    prods = [enrich_production(r) for r in rows]

    all_shorts = [enrich_short(s) for s in db_shorts_all()]

    summary = {
        'total':      len(prods),
        'iniciado':   sum(1 for p in prods if p['status'] == 'iniciado'),
        'andamento':  sum(1 for p in prods if p['status'] == 'em_andamento'),
        'pausado':    sum(1 for p in prods if p['status'] == 'pausado'),
        'concluido':  sum(1 for p in prods if p['status'] == 'concluido'),
        'atrasado':   sum(1 for p in prods if p['dl_color'] == 'red' and p['status'] != 'concluido'),
    }
    return render_template("productions.html",
                           prods=prods, summary=summary,
                           shorts=all_shorts,
                           active="productions")


@app.route("/productions/new", methods=["GET", "POST"])
def production_new():
    msg = msg_type = None
    tab = request.args.get("tab", "producao")
    if request.method == "POST":
        title      = request.form.get("title", "").strip()
        producer   = request.form.get("producer", "").strip()
        total_arcs = request.form.get("total_arcs", "11").strip()
        notes      = request.form.get("notes", "").strip()

        errors = []
        if not title:    errors.append("Título obrigatório")
        if not producer: errors.append("Produtor obrigatório")
        if not total_arcs.isdigit() or not (1 <= int(total_arcs) <= 15):
            errors.append("Número de arcos inválido (1–15)")

        if errors:
            msg = " | ".join(errors)
            msg_type = "error"
        else:
            now          = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            script_chars = int(request.form.get("script_chars", "0") or "0")
            arc_chars_raw = request.form.get("arc_chars", "")
            prod_type    = request.form.get("prod_type", "producao")
            if prod_type not in ("producao","roteiro","decupagem","edicao"):
                prod_type = "producao"
            try:
                json.loads(arc_chars_raw or "{}")
            except Exception:
                arc_chars_raw = ""
            with get_db() as c:
                c.execute(
                    ("INSERT INTO productions "
                     "(title,producer,prod_type,total_arcs,arcs_done,status,started_at,updated_at,notes,script_chars,arc_chars) "
                     "VALUES (?,?,?,?,'[]','iniciado',?,?,?,?,?)"),
                    (title, producer, prod_type, int(total_arcs), now, now, notes, script_chars, arc_chars_raw)
                )
            msg      = f"✔  Produção registrada: {title}"
            msg_type = "success"

    return render_template("production_new.html",
                           msg=msg, msg_type=msg_type,
                           tab=tab,
                           active="production_new")


@app.route("/productions/<int:pid>")
def production_detail(pid):
    with get_db() as c:
        p = c.execute("SELECT * FROM productions WHERE id=?", (pid,)).fetchone()
        if not p:
            return redirect(url_for("productions"))
        daily = c.execute(
            "SELECT * FROM production_daily WHERE production_id=? ORDER BY log_date DESC",
            (pid,)
        ).fetchall()
    prod = enrich_production(p)
    today = date.today().isoformat()
    return render_template("production_detail.html",
                           prod=prod,
                           daily=[dict(d) for d in daily],
                           today=today,
                           active="productions")

@app.route("/productions/<int:pid>/arc", methods=["POST"])
def production_arc(pid):
    arc_num = int(request.form.get("arc", 0))
    action  = request.form.get("action", "toggle")  # toggle | check | uncheck
    with get_db() as c:
        p = c.execute("SELECT * FROM productions WHERE id=?", (pid,)).fetchone()
        if not p:
            return redirect(url_for("productions"))
        done = json.loads(p["arcs_done"] or "[]")
        if action == "check" and arc_num not in done:
            done.append(arc_num)
        elif action == "uncheck" and arc_num in done:
            done.remove(arc_num)
        else:
            if arc_num in done:
                done.remove(arc_num)
            else:
                done.append(arc_num)
        done.sort()
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        # auto-complete if all arcs done
        status = p["status"]
        if len(done) >= p["total_arcs"] and status not in ("concluido",):
            status = "concluido"
            c.execute(
                "UPDATE productions SET arcs_done=?, status=?, updated_at=?, completed_at=? WHERE id=?",
                (json.dumps(done), status, now, now, pid)
            )
        else:
            if status == "concluido" and len(done) < p["total_arcs"]:
                status = "em_andamento"
            c.execute(
                "UPDATE productions SET arcs_done=?, updated_at=?, status=? WHERE id=?",
                (json.dumps(done), now, status, pid)
            )
    return redirect(url_for("production_detail", pid=pid))


@app.route("/productions/<int:pid>/status", methods=["POST"])
def production_status(pid):
    new_status = request.form.get("status", "")
    valid = ("iniciado", "em_andamento", "pausado", "concluido")
    if new_status not in valid:
        return redirect(url_for("production_detail", pid=pid))
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with get_db() as c:
        if new_status == "pausado":
            c.execute(
                "UPDATE productions SET status=?, paused_at=?, updated_at=? WHERE id=?",
                (new_status, now, now, pid)
            )
        elif new_status == "concluido":
            c.execute(
                "UPDATE productions SET status=?, completed_at=?, updated_at=? WHERE id=?",
                (new_status, now, now, pid)
            )
        else:
            c.execute(
                "UPDATE productions SET status=?, updated_at=? WHERE id=?",
                (new_status, now, pid)
            )
    return redirect(url_for("production_detail", pid=pid))




@app.route("/productions/<int:pid>/log", methods=["POST"])
def production_log(pid):
    chars    = request.form.get("chars", "0").strip()
    notes    = request.form.get("notes", "").strip()
    log_date = request.form.get("log_date", date.today().isoformat())
    if not chars.isdigit(): chars = "0"
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with get_db() as c:
        existing = c.execute(
            "SELECT id, chars_written FROM production_daily WHERE production_id=? AND log_date=?",
            (pid, log_date)
        ).fetchone()
        if existing:
            c.execute("UPDATE production_daily SET chars_written=?, notes=? WHERE id=?",
                      (existing["chars_written"] + int(chars), notes, existing["id"]))
        else:
            c.execute("INSERT INTO production_daily (production_id,log_date,chars_written,notes) VALUES (?,?,?,?)",
                      (pid, log_date, int(chars), notes))
        c.execute(
            "UPDATE productions SET updated_at=?, status=CASE WHEN status='iniciado' THEN 'em_andamento' ELSE status END WHERE id=?",
            (now_str, pid)
        )
    return redirect(url_for("production_detail", pid=pid))


@app.route("/productions/<int:pid>/log/<int:lid>/delete", methods=["POST"])
def log_delete(pid, lid):
    with get_db() as c:
        c.execute("DELETE FROM production_daily WHERE id=? AND production_id=?", (lid, pid))
    return redirect(url_for("production_detail", pid=pid))


@app.route("/productions/<int:pid>/delete", methods=["POST"])
@login_required
def production_delete(pid):
    with get_db() as c:
        c.execute("DELETE FROM production_daily WHERE production_id=?", (pid,))
        c.execute("DELETE FROM productions WHERE id=?", (pid,))
    return redirect(url_for("productions"))


@app.route("/reports/download")
@login_required
def reports_download():
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # ── Collect all data ──────────────────────────────────
    weeks_back = int(request.args.get("w", 0))
    rows, mon, sun = db_weekly(weeks_back)
    for r in rows:
        r["tag"] = ("high"   if r["count"] >= 5 else
                    "medium" if r["count"] >= 3 else "low")
    total_rr   = sum(r["count"] for r in rows)
    st         = db_stats()
    prods_by   = db_by_producer()
    prod_summ, all_prods, types_summ = db_production_report()
    shorts_summ, all_shorts_raw = db_shorts_report()
    all_shorts = [enrich_short(s) for s in all_shorts_raw]

    # ── Document setup ────────────────────────────────────
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Helpers ───────────────────────────────────────────
    def shade_para(para, hex_color):
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        pPr.append(shd)

    def shade_cell(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def section_title(text, hex_color):
        p = doc.add_paragraph()
        r = p.add_run(f'  {text}  ')
        r.font.name  = 'Courier New'
        r.font.size  = Pt(10)
        r.font.bold  = True
        r.font.color.rgb = RGBColor(0x08, 0x08, 0x10)
        shade_para(p, hex_color)
        doc.add_paragraph()

    def add_table(headers, col_widths, hdr_color_rgb):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = 'Table Grid'
        t.autofit = False
        for i, (h, w) in enumerate(zip(headers, col_widths)):
            c = t.rows[0].cells[i]
            c.width = w
            shade_cell(c, '111126')
            run = c.paragraphs[0].add_run(h)
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = hdr_color_rgb
        return t

    def add_row(table, values, col_widths, bg, fg_rgb):
        tr = table.add_row()
        for i, (v, w) in enumerate(zip(values, col_widths)):
            c = tr.cells[i]
            c.width = w
            shade_cell(c, bg)
            run = c.paragraphs[0].add_run(str(v))
            run.font.name  = 'Courier New'
            run.font.size  = Pt(9)
            run.font.color.rgb = fg_rgb
        return tr

    # ══════════════════════════════════════════════════════
    # PAGE 1 — HEADER
    # ══════════════════════════════════════════════════════
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_para(hdr, '0A0A1A')
    run = hdr.add_run('◆  PRODUÇÃO PHOENIX  ◆')
    run.font.name = 'Courier New'; run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x40, 0xC4, 0xFF)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_para(sub, '0A0A1A')
    r2 = sub.add_run('RELATÓRIO COMPLETO — REGRAVAÇÔES + PRODUÇÕES')
    r2.font.name = 'Courier New'; r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x88)

    period = doc.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_para(period, '0D0D1C')
    rp = period.add_run(
        f'Gerado em: {datetime.now().strftime("%d/%m/%Y %H:%M")}   |   '
        f'Período regravaçôes: {mon.strftime("%d/%m/%Y")} → {sun.strftime("%d/%m/%Y")}'
    )
    rp.font.name = 'Courier New'; rp.font.size = Pt(9)
    rp.font.color.rgb = RGBColor(0x40, 0xC4, 0xFF)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # SECTION 1 — RESUMO GERAL
    # ══════════════════════════════════════════════════════
    section_title('◆  RESUMO GERAL', '40C4FF')

    total_prods     = len(all_prods)
    total_concluido = sum(1 for p in all_prods if p['status'] == 'concluido')
    total_andamento = sum(1 for p in all_prods if p['status'] in ('em_andamento','iniciado'))

    summary_items = [
        ('REGRAVAÇÔES (SEMANA)', str(total_rr),      '1A0A0A' if total_rr>20 else '0A1A0A'),
        ('TOTAL REGRAVAÇÔES',    str(st['total']),   '0A0A1A'),
        ('PRODUÇÕES TOTAL',      str(total_prods),   '0A0A1A'),
        ('PRODUÇÕES CONCLUÍDAS', str(total_concluido),'0A1A0A'),
        ('EM ANDAMENTO',         str(total_andamento),'0A0A1A'),
    ]

    tbl = doc.add_table(rows=2, cols=len(summary_items))
    tbl.style = 'Table Grid'; tbl.autofit = False
    col_w = int(Inches(6.6) / len(summary_items))
    for i, (label, val, _) in enumerate(summary_items):
        ch = tbl.cell(0, i); cv = tbl.cell(1, i)
        ch.width = col_w; cv.width = col_w
        shade_cell(ch, '111126'); shade_cell(cv, '0A0A1A')
        rh = ch.paragraphs[0].add_run(label)
        rh.font.name='Courier New'; rh.font.size=Pt(7); rh.font.bold=True
        rh.font.color.rgb = RGBColor(0x40,0xC4,0xFF)
        ch.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        rv = cv.paragraphs[0].add_run(val)
        rv.font.name='Courier New'; rv.font.size=Pt(18); rv.font.bold=True
        rv.font.color.rgb = RGBColor(0x00,0xE6,0x76)
        cv.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # SECTION 2 — REGRAVAÇÔES DA SEMANA
    # ══════════════════════════════════════════════════════
    section_title('◆  REGRAVAÇÔES DA SEMANA', 'FFD700')

    if rows:
        cw = [Inches(1.3), Inches(1.4), Inches(0.5), Inches(1.9), Inches(0.9), Inches(0.8)]
        t = add_table(
            ['PRODUTOR','PROJETO','QTDE','MOTIVO','DATA','NÍVEL'],
            cw, RGBColor(0xFF,0xD7,0x00)
        )
        tag_map = {
            'high':   ('220808', RGBColor(0xFF,0xB0,0xBE), 'ALTO'),
            'medium': ('201800', RGBColor(0xFF,0xF0,0xA0), 'MÉDIO'),
            'low':    ('081808', RGBColor(0xB0,0xFF,0xD4), 'BAIXO'),
        }
        for row in rows:
            bg, fg, nivel = tag_map[row['tag']]
            add_row(t,
                [row['producer'], row['project'], row['count'],
                 row['reason'], row['created_at'][:10], nivel],
                cw, bg, fg)
    else:
        p = doc.add_paragraph('Nenhuma regravação registrada neste período.')
        p.runs[0].font.color.rgb = RGBColor(0x66,0x66,0x88)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # SECTION 3 — REGRAVAÇÔES POR PRODUTOR (acumulado)
    # ══════════════════════════════════════════════════════
    section_title('◆  REGRAVAÇÔES POR PRODUTOR — ACUMULADO', 'FF1744')

    if prods_by:
        cw2 = [Inches(2.3), Inches(1.0), Inches(3.0)]
        t2 = add_table(['PRODUTOR','TOTAL','MOTIVO PRINCIPAL'], cw2,
                        RGBColor(0xFF,0x17,0x44))
        for pb in prods_by:
            tv = pb['total']
            bg = ('220808' if tv>=10 else '201800' if tv>=5 else '081808')
            fg = (RGBColor(0xFF,0xB0,0xBE) if tv>=10 else
                  RGBColor(0xFF,0xF0,0xA0) if tv>=5 else
                  RGBColor(0xB0,0xFF,0xD4))
            add_row(t2, [pb['producer'], tv, pb.get('top_reason') or '—'],
                    cw2, bg, fg)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # SECTION 4 — RESUMO DE PRODUÇÕES POR PRODUTOR
    # ══════════════════════════════════════════════════════
    section_title('◆  PRODUÇÕES — RESUMO POR PRODUTOR', '00E676')

    if prod_summ:
        cw3 = [Inches(1.7), Inches(0.7), Inches(0.7), Inches(0.7),
               Inches(0.7), Inches(0.8), Inches(1.2)]
        t3 = add_table(
            ['PRODUTOR','TOTAL','INIC.','ANDAMENTO','PAUSADO','CONCLUÍDO','TEMPO MÉDIO'],
            cw3, RGBColor(0x00,0xE6,0x76)
        )
        for g in sorted(prod_summ, key=lambda x: x['total'], reverse=True):
            avg = f"{g['avg_bdays']} d.u." if g['avg_bdays'] is not None else 'sem concluídas'
            add_row(t3, [
                g['producer'], g['total'],
                g['iniciado'], g['em_andamento'],
                g['pausado'],  g['concluido'], avg
            ], cw3, '081808', RGBColor(0xB0,0xFF,0xD4))

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # SECTION 5 — HISTÓRICO COMPLETO DE PRODUÇÕES
    # ══════════════════════════════════════════════════════
    section_title('◆  PRODUÇÕES — HISTÓRICO COMPLETO', 'BB44FF')

    if all_prods:
        status_label = {
            'iniciado':    'INICIADO',
            'em_andamento':'EM ANDAMENTO',
            'pausado':     'PAUSADO',
            'concluido':   'CONCLUÍDO',
        }
        status_bg = {
            'iniciado':    ('0A0A1A', RGBColor(0xC0,0xEE,0xFF)),
            'em_andamento':('081808', RGBColor(0xB0,0xFF,0xD4)),
            'pausado':     ('201800', RGBColor(0xFF,0xF0,0xA0)),
            'concluido':   ('0A1A0A', RGBColor(0x80,0xFF,0xA0)),
        }
        cw4 = [Inches(1.5), Inches(2.2), Inches(0.8), Inches(1.2),
               Inches(0.8), Inches(0.6)]
        t4 = add_table(
            ['PRODUTOR','TÍTULO','ARCOS','STATUS','INÍCIO','DURAÇÃO'],
            cw4, RGBColor(0xBB,0x44,0xFF)
        )
        for p in sorted(all_prods, key=lambda x: x['started_at'], reverse=True):
            bg, fg = status_bg.get(p['status'], ('0A0A1A', RGBColor(0xF0,0xF0,0xF8)))
            arcos  = f"{p['arcs_done_count']}/{p['total_arcs']}"
            dur    = f"{p['duration_bdays']} d.u."
            add_row(t4, [
                p['producer'],
                p['title'][:32],
                arcos,
                status_label.get(p['status'], p['status']),
                p['started_at'][:10],
                dur,
            ], cw4, bg, fg)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # SECTION 6 — SHORTS — TEMPO MÉDIO POR PRODUTOR
    # ══════════════════════════════════════════════════════
    section_title('◆  SHORTS — TEMPO MÉDIO POR PRODUTOR', 'FFD700')

    if shorts_summ:
        cw5 = [Inches(1.8), Inches(0.7), Inches(0.8), Inches(0.7), Inches(0.8), Inches(1.1)]
        t5 = add_table(
            ['PRODUTOR','TOTAL','ANDAMENTO','PAUSADO','CONCLUÍDO','TEMPO MÉDIO'],
            cw5, RGBColor(0xFF,0xD7,0x00)
        )
        for g in sorted(shorts_summ, key=lambda x: x['total'], reverse=True):
            avg = f"{g['avg_bdays']} d.u." if g['avg_bdays'] is not None else 'sem concluídos'
            add_row(t5, [
                g['producer'], g['total'],
                g['andamento'], g['pausado'],
                g['concluido'], avg
            ], cw5, '1a1500', RGBColor(0xFF,0xF0,0xA0))
    else:
        p = doc.add_paragraph('Nenhum short registrado.')
        p.runs[0].font.color.rgb = RGBColor(0x66,0x66,0x88)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # SECTION 7 — SHORTS — HISTÓRICO COMPLETO
    # ══════════════════════════════════════════════════════
    section_title('◆  SHORTS — HISTÓRICO COMPLETO', 'BB44FF')

    if all_shorts:
        s_status = {
            'iniciado':    'INICIADO',
            'em_andamento':'EM ANDAMENTO',
            'pausado':     'PAUSADO',
            'concluido':   'CONCLUÍDO',
        }
        s_bg = {
            'iniciado':    ('0A0A1A', RGBColor(0xC0,0xEE,0xFF)),
            'em_andamento':('081808', RGBColor(0xB0,0xFF,0xD4)),
            'pausado':     ('201800', RGBColor(0xFF,0xF0,0xA0)),
            'concluido':   ('0A1A0A', RGBColor(0x80,0xFF,0xA0)),
        }
        cw6 = [Inches(1.5), Inches(2.2), Inches(1.2), Inches(0.9), Inches(0.9), Inches(0.7)]
        t6 = add_table(
            ['PRODUTOR','TÍTULO','STATUS','INÍCIO','CONCLUSÃO','DURAÇÃO'],
            cw6, RGBColor(0xBB,0x44,0xFF)
        )
        for s in sorted(all_shorts, key=lambda x: x['started_at'], reverse=True):
            bg, fg = s_bg.get(s['status'], ('0A0A1A', RGBColor(0xF0,0xF0,0xF8)))
            add_row(t6, [
                s['producer'],
                s['title'][:30],
                s_status.get(s['status'], s['status']),
                s['started_at'][:10],
                s['completed_at'][:10] if s.get('completed_at') else '—',
                f"{s['duration_bdays']} d.u.",
            ], cw6, bg, fg)
    else:
        p = doc.add_paragraph('Nenhum short registrado.')
        p.runs[0].font.color.rgb = RGBColor(0x66,0x66,0x88)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # FOOTER
    # ══════════════════════════════════════════════════════
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_para(footer_p, '0A0A1A')
    rf = footer_p.add_run(
        f'◆  Produção Phoenix  |  Relatório gerado automaticamente  |  '
        f'{datetime.now().strftime("%d/%m/%Y %H:%M")}  ◆'
    )
    rf.font.name = 'Courier New'; rf.font.size = Pt(8)
    rf.font.color.rgb = RGBColor(0x28, 0x28, 0x50)

    # ── Send file ─────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"relatorio_completo_{datetime.now().strftime('%d-%m-%Y_%H-%M')}.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
